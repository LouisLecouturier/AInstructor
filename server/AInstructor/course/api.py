import datetime
import os
import uuid as uuidLib

import aspose.words as pdf2md
import openai
import pdfplumber
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
from django.db import models
from django.shortcuts import get_object_or_404
from docx import Document
from ninja import Router, Schema, File, UploadedFile
from pydantic import Field
from django.http import HttpResponse


from app import models

router = Router(tags=["Course"])

"""________________________________________request conserning the courses__________________________________________________"""


class UploadTheme(Schema):
    theme: str = Field(...)
    name: str
    color: str

@router.post("/{user_id}")
def create_course(request, user_id: int, file : UploadedFile = File(...)):
    print(file.name)
    print(user_id)

    user = get_object_or_404(models.CustomUser, id=user_id)

    # Vérifier l'extension du fichier
    file_extension = os.path.splitext(file.name)[1].lower()

    if file_extension == '.pdf':
        # Convertir le fichier PDF en texte
        text_content = ""
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                text_content += page.extract_text()

    elif file_extension == '.docx':
        # Convertir le fichier Word en texte
        document = Document(file)
        paragraphs = document.paragraphs
        text_content = "\n".join([p.text for p in paragraphs])

    else:
        # Gérer d'autres extensions de fichiers ou afficher une erreur
        return {'error': 'Unsupported file format.'}

    # Enregistrer le contenu texte dans un fichier .txt
    txt_file_name = file.name.replace(file_extension, '.txt')
    txt_file_path = f"courses-for-IA/{txt_file_name}"  # Spécifiez le chemin de stockage souhaité

    # Convertir le texte en encodage UTF-8
    text_content_utf8 = text_content.encode('utf-8')

    txt_file = ContentFile(text_content_utf8)
    default_storage.save(txt_file_path, txt_file)

    # Sauvegarder le fichier .md
    course = models.Course.objects.create(
        name=file.name,
        uploadedBy=user,
        uploadedFile=file, 
        textPath=txt_file_path,
    )


    # Enregistrer le fichier .md en utilisant pdf2md.Document
    doc = pdf2md.Document(course.uploadedFile.path)
    doc.save(course.uploadedFile.path)

    print(course.uuid)
    return {'uuid': course.uuid}




@router.get("/teachers/{user_id}")
def get_my_courses(request, user_id: int):
    user = get_object_or_404(models.CustomUser, id=user_id)

    # teams = user.team_set.all()
    # courses = models.Course.objects.filter(team__in=teams)

    

    courses = models.Course.objects.filter(uploadedBy=user)

    Quizz = models.Quizz.objects.filter(course__in=courses)

    result = []
    for course in courses:
        quizz = models.Quizz.objects.filter(course=course)
        print(course.team.all().first())

        if course.team.all().first() is not None:
            teamName = course.team.all().first().name
        else :  teamName = None

        course_info = {
            'uuid': course.uuid,
            'name': course.name,
            'team': teamName,
            'description': course.description,
            'subject': course.subject,
            "status": "pending",
            "deliveryDate": course.deliveryDate,
            "creationDate": course.creationDate,
        }
        result.append(course_info)

    return result





   














































@router.get("/{uuid}/generate-questions", )
def generate_questions(request, uuid: str):
    """generate questions from the course"""
    course = get_object_or_404(models.Course, uuid=uuid)
    path = course.textPath
    # Ouvrez le fichier en mode lecture
    with open(path, "r", encoding="utf-8") as fichier:
        # Lisez le contenu du fichier
        texte = fichier.read()
    openai.api_key = "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"

    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "assistant", "content": texte},
            {"role": "user",
             "content": "Ecrit moi 10 questions sur ce texte pour tester mes connaisances mais tu ecris seulement les questions et pas les réponses"},
        ]
    )

    questions_with_numbers = response.choices[0].message.content.split('\n')
    questions = [q.split('.', 1)[1].strip() for q in questions_with_numbers if q.strip()]

    return {"questions": questions}


@router.get("byId/{uuid}")
def get_course_by_id(request, uuid: str):
    """get the course by id"""
    course = get_object_or_404(models.Course, uuid=uuid)
    return {
        'uuid': course.uuid,
        'name': course.name,
        'subject': course.subject,
        'text': course.text,
        'uploadedBy': course.uploadedBy.username,
        'color': course.color,
        'file': course.uploadedFile.name,
    }







class AssignCourse(Schema):
    uuid: uuidLib.UUID = Field(...)
    teamUUID: uuidLib.UUID = Field(...)
    deadline: datetime.date = Field(...)


@router.post("/assign-course")
def assign_course(request, body: AssignCourse):
    """assigne a course to a team"""
    course = get_object_or_404(models.Course, uuid=body.uuid)
    team = get_object_or_404(models.Team, uuid=body.teamUUID)
    course.team.add(team)
    course.dateEnd = body.deadline
    course.save()
    return {'uuid': course.uuid, 'name': course.name, 'teamName': team.name, 'teamUUID': team.uuid,
            'deadline': course.dateEnd}


# à mettre dans Team api
@router.get("/team/{uuid}")
def get_courses_by_team(request, uuid: uuidLib.UUID):
    """get all the courses of the team"""
    team = get_object_or_404(models.Team, uuid=uuid)
    courses = models.Course.objects.filter(team=team)

    result = []
    for course in courses:
        course_info = {
            'uuid': course.uuid,
            'name': course.name,
            'theme': course.theme,
            'uploadedBy': course.uploadedBy.username,
            'color': course.color,
            'file': course.uploaded_file.name,
            # 'text': course.text,
        }
        result.append(course_info)
    return result


class AssignCourse(Schema):
    course_id: uuidLib.UUID = Field(...)
    group_id: str = Field(...)
    deadline: datetime.date = Field(...)


@router.post("/assign-course")
def assign_course(request, body: AssignCourse):
    """assigne a course to a group"""
    course = get_object_or_404(models.Course, course_id=body.course_id)
    group = get_object_or_404(models.Groupe, group_id=body.group_id)
    course.group.add(group)
    course.date_end = body.deadline
    course.save()
    return {'course_id': course.course_id, 'course name': course.name, 'group name': group.name,
            'group_id': group.group_id, 'deadline': course.date_end}


@router.get("/courses/{group_id}")
def get_courses_by_group(request, group_id: uuidLib.UUID):
    """get all the courses of the user"""
    group = get_object_or_404(models.Groupe, group_id=group_id)
    courses = models.Course.objects.filter(group=group)

    result = []
    for course in courses:
        course_info = {
            'course_id': course.course_id,
            'name': course.name,
            'theme': course.theme,
            'uploaded_by': course.uploaded_by.username,
            'color': course.color,
            'file': course.uploaded_file.name,
            # 'text': course.text,
        }
        result.append(course_info)
    return result











class UpdateCourse(Schema):
    name: str = Field(...)
    subject: str = Field(...)
    # color: str = Field(...)
    description: str = Field(...)


@router.put("/put/{uuid}")
def update_meta_data_from_course(request, uuid,  courseInfo: UpdateCourse):
    """update the course metadata : name, theme, color"""
    course = get_object_or_404(models.Course, uuid=uuid)
    print(courseInfo)

    course.subject = courseInfo.subject
    course.name = courseInfo.name
    course.description = courseInfo.description
    # course.color = courseInfo.color
    course.save()

    return {
        'uuid': course.uuid, 
        'subject': course.subject, 
        'name': course.name,
        'description': course.description,
        # 'color': course.color
        }















class UpdateCourseFile(Schema):
    uuid: uuidLib.UUID = Field(...)


@router.put("/update-course-file")
def update_course_file(request, body: UpdateCourseFile, file: UploadedFile = File(...)):
    """update the course file"""
    course = get_object_or_404(models.Course, uuid=body.uuid)
    if file.name.endswith(('.pdf', '.md')):
        course.uploadedFile = file
        course.save()
        return {'uuid': str(course.uuid)}
    else:
        return {'error': 'File is not a PDF or MD.'}


@router.delete("/course/{uuid}")
def delete_data(request, uuid: str):
    course = models.Course.objects.get(uuid=uuid)
    course.delete()
    return {'uuid': uuid}


class UpdateCourseText(Schema):
    uuid: uuidLib.UUID = Field(...)
    text: str = Field(...)


@router.put("/update-course-text")
def update_course_text(request, body: UpdateCourseText):
    """update the course text"""
    course = get_object_or_404(models.Course, uuid=body.uuid)
    course.text = body.text
    course.save()
    return {'uuid': course.uuid, 'text': course.text}
